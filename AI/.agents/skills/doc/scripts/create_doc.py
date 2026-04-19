#!/usr/bin/env python3
"""Generate a DOCX file from structured JSON content."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUTPUT_DIR = Path("output/doc")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create a .docx file from structured title/section content."
    )
    parser.add_argument(
        "--input-json",
        help="Path to a JSON file containing title and sections.",
    )
    parser.add_argument(
        "--content-json",
        help="Inline JSON string containing title and sections.",
    )
    parser.add_argument(
        "--output",
        help="Optional output path. Defaults to output/doc/<slug>.docx",
    )
    return parser.parse_args()


def load_payload(args: argparse.Namespace) -> dict[str, Any]:
    if bool(args.input_json) == bool(args.content_json):
        raise SystemExit("Provide exactly one of --input-json or --content-json.")

    if args.input_json:
        payload = json.loads(Path(args.input_json).read_text(encoding="utf-8"))
    else:
        payload = json.loads(args.content_json)

    if not isinstance(payload, dict):
        raise SystemExit("Document payload must be a JSON object.")

    title = payload.get("title")
    sections = payload.get("sections")
    if not isinstance(title, str) or not title.strip():
        raise SystemExit("Payload must include a non-empty string `title`.")
    if not isinstance(sections, list) or not sections:
        raise SystemExit("Payload must include a non-empty list `sections`.")
    return payload


def slugify(value: str) -> str:
    chars = []
    for char in value.lower():
        if char.isalnum():
            chars.append(char)
        elif chars and chars[-1] != "-":
            chars.append("-")
    return "".join(chars).strip("-") or "document"


def normalize_paragraphs(section: dict[str, Any]) -> list[str]:
    body = section.get("body")
    paragraphs = section.get("paragraphs")
    if isinstance(paragraphs, list) and paragraphs:
        return [str(item).strip() for item in paragraphs if str(item).strip()]
    if isinstance(body, str) and body.strip():
        return [part.strip() for part in body.split("\n\n") if part.strip()]
    raise SystemExit("Each section needs `paragraphs` or `body` content.")


def set_default_font(document: Document) -> None:
    styles = document.styles
    normal_style = styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(12)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)


def build_document(payload: dict[str, Any]) -> Document:
    document = Document()
    set_default_font(document)
    configure_page(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(payload["title"].strip())
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(16)

    for index, raw_section in enumerate(payload["sections"], start=1):
        if not isinstance(raw_section, dict):
            raise SystemExit("Each section must be an object.")
        heading = str(raw_section.get("heading", "")).strip()
        if not heading:
            raise SystemExit("Each section must include a non-empty `heading`.")
        document.add_heading(f"{index}. {heading}", level=1)
        for paragraph_text in normalize_paragraphs(raw_section):
            paragraph = document.add_paragraph(paragraph_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Inches(0.3)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(6)

    return document


def resolve_output_path(title: str, output: str | None) -> Path:
    if output:
        return Path(output)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR / f"{slugify(title)}.docx"


def main() -> int:
    args = parse_args()
    payload = load_payload(args)
    document = build_document(payload)
    output_path = resolve_output_path(payload["title"], args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    print(output_path.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
